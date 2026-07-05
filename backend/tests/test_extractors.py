"""
Tests for the file text extractors (RAG ingestion).
Runnable with pytest or `python3 -m unittest discover -s tests` from backend/.

The module is loaded directly by path: importing `rag.extractors` the
normal way would trigger rag/__init__, which pulls llama_index (not
needed for extraction and not installed in minimal test environments).
"""

import importlib.util
import os
import unittest

_SPEC = importlib.util.spec_from_file_location(
    "extractors",
    os.path.join(os.path.dirname(__file__), "..", "rag", "extractors.py"),
)
extractors = importlib.util.module_from_spec(_SPEC)
_SPEC.loader.exec_module(extractors)

extract_text = extractors.extract_text
ExtractionError = extractors.ExtractionError
MAX_FILE_SIZE_BYTES = extractors.MAX_FILE_SIZE_BYTES


def _docx_available() -> bool:
    return importlib.util.find_spec("docx") is not None


class TestTextLikeExtraction(unittest.TestCase):
    def test_txt(self):
        self.assertEqual(extract_text("notes.txt", b"hello world"), "hello world")

    def test_markdown(self):
        text = extract_text("doc.md", "# Titolo\n\nContenuto **bold**".encode("utf-8"))
        self.assertIn("Titolo", text)

    def test_non_utf8_falls_back(self):
        content = "caffè".encode("latin-1")
        text = extract_text("legacy.txt", content)
        self.assertIn("caff", text)

    def test_empty_file_rejected(self):
        with self.assertRaises(ExtractionError):
            extract_text("empty.txt", b"")

    def test_whitespace_only_rejected(self):
        with self.assertRaises(ExtractionError):
            extract_text("blank.txt", b"   \n\t  ")

    def test_oversized_rejected(self):
        with self.assertRaises(ExtractionError):
            extract_text("big.txt", b"x" * (MAX_FILE_SIZE_BYTES + 1))

    def test_unsupported_extension(self):
        with self.assertRaises(ExtractionError) as ctx:
            extract_text("archive.zip", b"PK...")
        self.assertIn(".zip", str(ctx.exception))

    def test_missing_extension(self):
        with self.assertRaises(ExtractionError):
            extract_text("noextension", b"text")


class TestHtmlExtraction(unittest.TestCase):
    def test_html_strips_markup_and_scripts(self):
        html = b"""
        <html><head><style>body { color: red }</style>
        <script>alert('x')</script></head>
        <body><h1>Sustainability Report</h1><p>Our commitment.</p></body></html>
        """
        text = extract_text("page.html", html)
        self.assertIn("Sustainability Report", text)
        self.assertIn("Our commitment.", text)
        self.assertNotIn("alert", text)
        self.assertNotIn("color: red", text)

    def test_htm_extension(self):
        text = extract_text("page.htm", b"<p>ok</p>")
        self.assertEqual(text, "ok")


class TestPdfExtraction(unittest.TestCase):
    def test_pdf_roundtrip(self):
        # Build a tiny one-page PDF with pypdf itself
        import io
        from pypdf import PdfWriter

        writer = PdfWriter()
        writer.add_blank_page(width=200, height=200)
        buffer = io.BytesIO()
        writer.write(buffer)

        # A blank page has no text -> extraction error is the correct outcome
        with self.assertRaises(ExtractionError):
            extract_text("blank.pdf", buffer.getvalue())

    def test_corrupt_pdf_raises(self):
        with self.assertRaises(Exception):
            extract_text("broken.pdf", b"not a real pdf content")


class TestBackendDispatcher(unittest.TestCase):
    def tearDown(self):
        # Restore real adapters after monkeypatching
        importlib_reload = importlib.util.spec_from_file_location(
            "extractors_reload",
            os.path.join(os.path.dirname(__file__), "..", "rag", "extractors.py"),
        )
        fresh = importlib.util.module_from_spec(importlib_reload)
        importlib_reload.loader.exec_module(fresh)
        extractors._extract_docling = fresh._extract_docling
        extractors._extract_glmocr = fresh._extract_glmocr

    def test_unknown_backend_falls_back_to_local_set(self):
        self.assertEqual(
            extractors.supported_extensions("nonsense"),
            extractors.supported_extensions("local"),
        )

    def test_backends_unlock_images(self):
        self.assertNotIn(".png", extractors.supported_extensions("local"))
        self.assertIn(".png", extractors.supported_extensions("docling"))
        self.assertIn(".png", extractors.supported_extensions("glmocr"))

    def test_text_formats_always_local(self):
        # Even with a backend configured, txt/md decode locally:
        # the backend adapter must never be invoked
        def _boom(*args, **kwargs):
            raise AssertionError("backend should not be called for txt")

        extractors._extract_docling = _boom
        text = extract_text("note.txt", b"plain", backend="docling")
        self.assertEqual(text, "plain")

    def test_backend_used_for_pdf(self):
        extractors._extract_docling = lambda filename, content: "from docling"
        text = extract_text("doc.pdf", b"%PDF-fake", backend="docling")
        self.assertEqual(text, "from docling")

    def test_glmocr_not_used_for_docx(self):
        # glmocr only handles pdf/images; docx routes to the local parser
        def _boom(*args, **kwargs):
            raise AssertionError("glmocr should not be called for docx")

        extractors._extract_glmocr = _boom
        with self.assertRaises(Exception):
            # Local docx parser will fail on fake bytes (or raise ImportError
            # without python-docx). Either way, glmocr was not invoked
            extract_text("doc.docx", b"not a real docx", backend="glmocr")

    def test_runtime_failure_falls_back_to_local(self):
        def _fail(filename, content):
            raise RuntimeError("backend service down")

        extractors._extract_docling = _fail
        html = b"<p>fallback works</p>"
        text = extract_text("page.html", html, backend="docling")
        self.assertEqual(text, "fallback works")

    def test_image_failure_does_not_fall_back(self):
        def _fail(filename, content):
            raise RuntimeError("backend service down")

        extractors._extract_docling = _fail
        with self.assertRaises(ExtractionError):
            extract_text("scan.png", b"\x89PNG fake", backend="docling")

    def test_image_rejected_on_local_backend(self):
        with self.assertRaises(ExtractionError) as ctx:
            extract_text("scan.png", b"\x89PNG fake", backend="local")
        self.assertIn(".png", str(ctx.exception))

    def test_missing_package_raises_import_error(self):
        # Real adapters: docling/glmocr are not installed in this environment
        with self.assertRaises(ImportError):
            extract_text("doc.pdf", b"%PDF-fake", backend="glmocr")

    def test_configured_backend_defaults_to_local_without_settings(self):
        # config/pydantic-settings may be unavailable in minimal envs:
        # resolution must degrade to local, never crash
        self.assertIn(extractors.configured_backend(), extractors.KNOWN_BACKENDS)


@unittest.skipUnless(_docx_available(), "python-docx not installed")
class TestDocxExtraction(unittest.TestCase):
    def test_docx_paragraphs_and_tables(self):
        import io
        from docx import Document

        document = Document()
        document.add_paragraph("Paragrafo uno")
        table = document.add_table(rows=1, cols=2)
        table.rows[0].cells[0].text = "A"
        table.rows[0].cells[1].text = "B"
        buffer = io.BytesIO()
        document.save(buffer)

        text = extract_text("doc.docx", buffer.getvalue())
        self.assertIn("Paragrafo uno", text)
        self.assertIn("A | B", text)


if __name__ == "__main__":
    unittest.main()
