"""
Text Extractors
Turn uploaded files into plain text for the RAG pipeline.

The extraction backend is configuration, not code (EXTRACTOR_BACKEND):

- "local" (default): pypdf / python-docx / BeautifulSoup. Zero extra
  dependencies, instant, and documents never leave the machine. Weak on
  complex PDF layouts and unable to OCR scans.
- "docling": IBM Docling, still fully local, no GPU required. Better
  table/layout extraction; also unlocks image files. Heavy dependencies
  (PyTorch). `pip install docling`
- "glmocr": GLM-OCR (0.9B vision model, MIT). State-of-the-art document
  parsing incl. scanned documents; self-hosted (vLLM/Ollama, ~2-4GB
  VRAM) keeps data in-house, or via the Z.ai cloud API (data leaves
  your infra — opt-in consciously). `pip install glmocr`

Routing rules:
- txt/md are always decoded locally (no backend adds value there).
- A backend only handles the formats it is good at (see _BACKEND_FORMATS);
  everything else falls through to the local parsers.
- Runtime failures of a configured backend fall back to local with a
  warning (fail-open) — except formats only the backend can handle
  (e.g. images), which fail explicitly.
- A configured backend whose package is missing raises ImportError
  loudly: that is a deployment mistake, not a degradation to hide.
"""

import io
import logging
import os
import tempfile
from typing import Optional, Set

logger = logging.getLogger(__name__)

MAX_FILE_SIZE_BYTES = 10 * 1024 * 1024  # 10 MB

BACKEND_LOCAL = "local"
BACKEND_DOCLING = "docling"
BACKEND_GLMOCR = "glmocr"
KNOWN_BACKENDS = (BACKEND_LOCAL, BACKEND_DOCLING, BACKEND_GLMOCR)

_TEXT_EXTENSIONS = {".txt", ".md", ".markdown"}
_LOCAL_EXTENSIONS = {".pdf", ".docx", ".html", ".htm"} | _TEXT_EXTENSIONS
_IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".webp", ".tiff"}

# Formats each backend is actually good at; anything else goes local
_BACKEND_FORMATS = {
    BACKEND_DOCLING: {".pdf", ".docx", ".html", ".htm"} | _IMAGE_EXTENSIONS,
    BACKEND_GLMOCR: {".pdf"} | _IMAGE_EXTENSIONS,
}


class ExtractionError(ValueError):
    """File could not be turned into text (unsupported, empty, corrupt)."""


def configured_backend() -> str:
    """Resolve the extractor backend from settings (default: local)."""
    try:
        from config import settings

        backend = (settings.extractor_backend or BACKEND_LOCAL).strip().lower()
    except Exception:
        return BACKEND_LOCAL

    if backend not in KNOWN_BACKENDS:
        logger.warning("Unknown EXTRACTOR_BACKEND %r, using local", backend)
        return BACKEND_LOCAL
    return backend


def supported_extensions(backend: Optional[str] = None) -> Set[str]:
    """File extensions accepted with the given (or configured) backend."""
    backend = backend or configured_backend()
    return _LOCAL_EXTENSIONS | _BACKEND_FORMATS.get(backend, set())


def extract_text(
    filename: str,
    content: bytes,
    backend: Optional[str] = None,
) -> str:
    """
    Extract plain text from an uploaded file.

    Args:
        filename: Original file name (the extension selects the parser).
        content: Raw file bytes.
        backend: Override the configured EXTRACTOR_BACKEND.

    Returns:
        Extracted text, stripped.

    Raises:
        ExtractionError: unsupported type, oversized, empty result.
        ImportError: configured backend's package not installed.
    """
    if not content:
        raise ExtractionError("Empty file")
    if len(content) > MAX_FILE_SIZE_BYTES:
        raise ExtractionError(
            f"File too large ({len(content)} bytes, max {MAX_FILE_SIZE_BYTES})"
        )

    backend = (backend or configured_backend()).lower()
    extension = os.path.splitext(filename or "")[1].lower()

    allowed = supported_extensions(backend)
    if extension not in allowed:
        raise ExtractionError(
            f"Unsupported file type {extension or '(none)'} for backend "
            f"'{backend}'; supported: {', '.join(sorted(allowed))}"
        )

    # Plain text formats: no backend beats decoding the bytes
    if extension in _TEXT_EXTENSIONS:
        return _finalize(_decode_text(content))

    backend_handles = extension in _BACKEND_FORMATS.get(backend, set())

    if backend != BACKEND_LOCAL and backend_handles:
        try:
            return _finalize(_extract_with_backend(backend, filename, content, extension))
        except ImportError:
            # Deployment mistake: surface it, don't silently degrade quality
            raise
        except Exception as e:
            if extension in _IMAGE_EXTENSIONS:
                # Local parsers cannot handle images: nothing to fall back to
                raise ExtractionError(
                    f"{backend} extraction failed for image: {e}"
                ) from e
            logger.warning(
                "%s extraction failed (%s); falling back to local parsers",
                backend, e,
            )

    return _finalize(_extract_local(extension, content))


def _finalize(text: str) -> str:
    text = (text or "").strip()
    if not text:
        raise ExtractionError("No text could be extracted from the file")
    return text



# Local backend (pypdf / python-docx / BeautifulSoup) 
def _extract_local(extension: str, content: bytes) -> str:
    if extension == ".pdf":
        return _extract_pdf(content)
    if extension == ".docx":
        return _extract_docx(content)
    if extension in (".html", ".htm"):
        return _extract_html(content)
    if extension in _IMAGE_EXTENSIONS:
        raise ExtractionError(
            "Image files require EXTRACTOR_BACKEND=docling or glmocr"
        )
    return _decode_text(content)


def _decode_text(content: bytes) -> str:
    try:
        return content.decode("utf-8")
    except UnicodeDecodeError:
        return content.decode("latin-1", errors="replace")


def _extract_pdf(content: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as e:
        raise ImportError("PDF extraction requires 'pypdf' (pip install pypdf)") from e

    reader = PdfReader(io.BytesIO(content))
    pages = []
    for page in reader.pages:
        try:
            pages.append(page.extract_text() or "")
        except Exception as e:
            logger.warning("PDF page extraction failed: %s", e)
    return "\n\n".join(pages)


def _extract_docx(content: bytes) -> str:
    try:
        from docx import Document
    except ImportError as e:
        raise ImportError(
            "DOCX extraction requires 'python-docx' (pip install python-docx)"
        ) from e

    document = Document(io.BytesIO(content))
    parts = [p.text for p in document.paragraphs if p.text and p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _extract_html(content: bytes) -> str:
    try:
        from bs4 import BeautifulSoup
    except ImportError as e:
        raise ImportError(
            "HTML extraction requires 'beautifulsoup4' (pip install beautifulsoup4)"
        ) from e

    soup = BeautifulSoup(_decode_text(content), "html.parser")
    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()
    return soup.get_text(separator="\n")



# Optional backends
def _extract_with_backend(
    backend: str,
    filename: str,
    content: bytes,
    extension: str,
) -> str:
    if backend == BACKEND_DOCLING:
        return _extract_docling(filename, content)
    if backend == BACKEND_GLMOCR:
        return _extract_glmocr(content, extension)
    raise ExtractionError(f"Unknown extractor backend: {backend}")


# Docling loads layout models at startup: keep one converter per process
_docling_converter = None


def _extract_docling(filename: str, content: bytes) -> str:
    try:
        from docling.datamodel.base_models import DocumentStream
        from docling.document_converter import DocumentConverter
    except ImportError as e:
        raise ImportError(
            "EXTRACTOR_BACKEND=docling requires 'docling' (pip install docling)"
        ) from e

    global _docling_converter
    if _docling_converter is None:
        _docling_converter = DocumentConverter()

    stream = DocumentStream(name=filename or "document", stream=io.BytesIO(content))
    result = _docling_converter.convert(stream)
    return result.document.export_to_markdown()


def _extract_glmocr(content: bytes, extension: str) -> str:
    try:
        import glmocr
    except ImportError as e:
        raise ImportError(
            "EXTRACTOR_BACKEND=glmocr requires 'glmocr' "
            "(pip install glmocr, or glmocr[selfhosted] for vLLM/Ollama)"
        ) from e

    # Forward optional settings to the SDK without overriding an
    # environment the operator configured directly
    try:
        from config import settings

        if getattr(settings, "glmocr_api_key", None):
            os.environ.setdefault("GLMOCR_API_KEY", settings.glmocr_api_key)
        if getattr(settings, "glmocr_base_url", None):
            os.environ.setdefault("GLMOCR_BASE_URL", settings.glmocr_base_url)
    except Exception:
        pass

    # The SDK works on file paths
    with tempfile.NamedTemporaryFile(suffix=extension, delete=False) as handle:
        handle.write(content)
        tmp_path = handle.name

    try:
        result = glmocr.parse(tmp_path)
    finally:
        try:
            os.unlink(tmp_path)
        except OSError:
            pass

    # Be liberal in what we accept from the SDK result object
    for attribute in ("markdown", "text", "content"):
        value = getattr(result, attribute, None)
        if isinstance(value, str) and value.strip():
            return value
    for method in ("to_markdown", "to_text"):
        candidate = getattr(result, method, None)
        if callable(candidate):
            value = candidate()
            if isinstance(value, str) and value.strip():
                return value
    return str(result)
